"""File importer for GPGNotes - converts various formats to markdown notes."""

import re
from pathlib import Path
from typing import Optional


class ImportError(Exception):
    """Raised when file import fails."""

    pass


class MissingDependencyError(ImportError):
    """Raised when required dependency is not installed."""

    pass


def _check_dependency(module_name: str, package_name: str):
    """Check if a dependency is available, raise helpful error if not."""
    try:
        __import__(module_name)
    except ModuleNotFoundError:
        raise MissingDependencyError(
            f"The '{package_name}' package is required to import this file type.\n"
            f"Install it with: pip install gpgnotes[import]\n"
            f"Or: pip install {package_name}"
        )


def import_markdown(file_path: Path) -> tuple[str, str]:
    """
    Import a markdown file.

    Args:
        file_path: Path to the markdown file

    Returns:
        Tuple of (title, content)
    """
    content = file_path.read_text(encoding="utf-8")

    # Try to extract title from first heading
    title = file_path.stem
    lines = content.split("\n")
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            break

    return title, content


def import_text(file_path: Path) -> tuple[str, str]:
    """
    Import a plain text file.

    Args:
        file_path: Path to the text file

    Returns:
        Tuple of (title, content)
    """
    content = file_path.read_text(encoding="utf-8")
    title = file_path.stem

    # Use first non-empty line as title if it looks like a title
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if line and len(line) < 100:
            title = line
            break

    return title, content


def import_rtf(file_path: Path) -> tuple[str, str]:
    """
    Import an RTF file, converting to plain text.

    Args:
        file_path: Path to the RTF file

    Returns:
        Tuple of (title, content)
    """
    _check_dependency("striprtf", "striprtf")
    from striprtf.striprtf import rtf_to_text

    rtf_content = file_path.read_text(encoding="utf-8", errors="ignore")
    content = rtf_to_text(rtf_content)

    # Clean up the content
    content = content.strip()

    # Use filename as title, or first line if short enough
    title = file_path.stem
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if line and len(line) < 100:
            title = line
            break

    return title, content


def import_pdf(file_path: Path) -> tuple[str, str]:
    """
    Import a PDF file, extracting text content.

    Args:
        file_path: Path to the PDF file

    Returns:
        Tuple of (title, content)
    """
    _check_dependency("pypdf", "pypdf")
    from pypdf import PdfReader

    reader = PdfReader(file_path)

    # Extract text from all pages
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    content = "\n\n".join(text_parts)

    # Try to get title from PDF metadata
    title = file_path.stem
    if reader.metadata and reader.metadata.title:
        title = reader.metadata.title

    # Clean up content - normalize whitespace
    content = re.sub(r"\n{3,}", "\n\n", content)
    content = content.strip()

    return title, content


def import_docx(file_path: Path) -> tuple[str, str]:
    """
    Import a Word document (.docx), converting to markdown.

    Args:
        file_path: Path to the docx file

    Returns:
        Tuple of (title, content)
    """
    _check_dependency("docx", "python-docx")
    from docx import Document

    doc = Document(file_path)

    # Extract content with basic formatting preservation
    content_parts = []
    title = file_path.stem

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            content_parts.append("")
            continue

        # Check for headings
        if para.style.name.startswith("Heading"):
            level = 1
            try:
                level = int(para.style.name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            prefix = "#" * level
            content_parts.append(f"{prefix} {text}")

            # Use first heading as title
            if i == 0 or (title == file_path.stem and level == 1):
                title = text
        else:
            # Handle basic inline formatting
            formatted_text = _format_docx_paragraph(para)
            content_parts.append(formatted_text)

    # Handle tables
    for table in doc.tables:
        content_parts.append("")
        content_parts.append(_convert_docx_table(table))

    content = "\n\n".join(content_parts)

    # Clean up multiple blank lines
    content = re.sub(r"\n{3,}", "\n\n", content)
    content = content.strip()

    return title, content


def _format_docx_paragraph(para) -> str:
    """Format a docx paragraph with basic markdown formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue

        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        parts.append(text)

    return "".join(parts)


def _convert_docx_table(table) -> str:
    """Convert a docx table to markdown format."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

        # Add header separator after first row
        if i == 0:
            separator = "| " + " | ".join(["---"] * len(cells)) + " |"
            rows.append(separator)

    return "\n".join(rows)


def import_file(file_path: Path, title: Optional[str] = None) -> tuple[str, str]:
    """
    Import a file based on its extension.

    Args:
        file_path: Path to the file to import
        title: Optional custom title (overrides auto-detected title)

    Returns:
        Tuple of (title, content)

    Raises:
        ImportError: If file type is not supported
        FileNotFoundError: If file does not exist
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    importers = {
        ".md": import_markdown,
        ".markdown": import_markdown,
        ".txt": import_text,
        ".text": import_text,
        ".rtf": import_rtf,
        ".pdf": import_pdf,
        ".docx": import_docx,
    }

    if suffix not in importers:
        supported = ", ".join(sorted(importers.keys()))
        raise ImportError(f"Unsupported file type: {suffix}\nSupported formats: {supported}")

    detected_title, content = importers[suffix](file_path)

    # Use custom title if provided
    final_title = title if title else detected_title

    return final_title, content


def get_supported_extensions() -> list[str]:
    """Return list of supported file extensions."""
    return [".md", ".markdown", ".txt", ".text", ".rtf", ".pdf", ".docx"]
